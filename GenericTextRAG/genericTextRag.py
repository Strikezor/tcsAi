import streamlit as st
import tempfile
import os
import httpx
from dotenv import load_dotenv

# Document Loaders & Extractors
from pdfminer.high_level import extract_text as extract_pdf_text
import docx

# LangChain Imports
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_classic.chains import create_retrieval_chain
from langchain_classic.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.documents import Document

# Load environment variables
load_dotenv()

# --- Configuration & Setup ---
tiktoken_cache_dir = "./token"
os.environ["TIKTOKEN_CACHE_DIR"] = tiktoken_cache_dir

# Client setup to bypass SSL verification (specific to your environment)
client = httpx.Client(verify=False)

# Page Config
st.set_page_config(page_title="Multi-Format Doc Chatbot", layout="centered")
st.title("🤖 Chat with your Documents")

# --- Helper Functions ---

def get_text_from_file(file_path, file_extension):
    """
    Extract text based on file extension.
    Supported: .pdf, .docx, .txt
    """
    text = ""
    try:
        if file_extension == ".pdf":
            text = extract_pdf_text(file_path)
        elif file_extension == ".docx":
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        elif file_extension == ".txt":
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        elif file_extension == ".doc":
            st.warning("⚠️ Legacy .doc files are not fully supported. Please convert to .docx or .pdf for best results.")
            return ""
    except Exception as e:
        st.error(f"Error extracting text: {e}")
    return text

def initialize_vector_db(uploaded_file):
    """
    Process the file, chunk it, and create/update the Vector DB.
    Returns the retrieval chain.
    """
    # Create a temp file to process
    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
        temp_file.write(uploaded_file.read())
        temp_file_path = temp_file.name

    try:
        # 1. Extract Text
        raw_text = get_text_from_file(temp_file_path, file_extension)
        
        if not raw_text:
            st.error("Could not extract text from the file.")
            return None

        # 2. Chunking
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000, 
            chunk_overlap=200,
            length_function=len,
            is_separator_regex=False
        )
        # We wrap the text in a Document object for LangChain compatibility
        chunks = text_splitter.split_text(raw_text)
        documents = [Document(page_content=chunk) for chunk in chunks]

        # 3. Embedding Model
        embedding_model = OpenAIEmbeddings( 
            base_url="https://genailab.tcs.in", 
            model="azure/genailab-maas-text-embedding-3-large", 
            api_key=os.getenv("OPENAI_API_KEY"),
            http_client=client
        ) 

        # 4. Store in Chroma (Using ephemeral storage for session)
        # Note: We use a new collection or clear old one if needed. 
        # For simplicity here, we create a new in-memory instance or persist to a temp folder.
        vectordb = Chroma.from_documents(
            documents, 
            embedding_model,
            # persist_directory="./chroma_db_session" # Optional: Persist to disk
        )

        # 5. LLM Setup
        llm = ChatOpenAI( 
            base_url="https://genailab.tcs.in", 
            model="azure_ai/genailab-maas-DeepSeek-V3-0324", 
            api_key=os.getenv("OPENAI_API_KEY"),
            http_client=client 
        )

        # 6. Create Chain
        retriever = vectordb.as_retriever(search_type="similarity", search_kwargs={"k": 5})
        
        prompt = ChatPromptTemplate.from_template("""
        You are a factual assistant designed to answer questions strictly based on the provided documents.
        
        Instructions:
        1. Use ONLY the context provided below to answer the question.
        2. If the answer is not explicitly present in the context, say "I cannot find the answer in the provided document."
        3. Do not make up information, assume facts, or use outside knowledge.
        4. Keep your answer concise and relevant to the context.

        <context>
        {context}
        </context>
        
        Question: {input}
        """)

        combine_docs_chain = create_stuff_documents_chain(llm, prompt)
        rag_chain = create_retrieval_chain(retriever, combine_docs_chain)
        
        return rag_chain

    finally:
        # Cleanup temp file
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)

# --- Main UI Logic ---

# Initialize Session State
if "messages" not in st.session_state:
    st.session_state.messages = []

if "rag_chain" not in st.session_state:
    st.session_state.rag_chain = None

if "current_file" not in st.session_state:
    st.session_state.current_file = None

# Sidebar for Upload
with st.sidebar:
    st.header("📁 Document Upload")
    upload_file = st.file_uploader(
        "Upload a document", 
        type=["pdf", "docx", "doc", "txt"]
    )
    
    # Process file only if it's new
    if upload_file:
        file_id = upload_file.name + str(upload_file.size)
        if st.session_state.current_file != file_id:
            with st.spinner("Processing document... This may take a moment."):
                chain = initialize_vector_db(upload_file)
                if chain:
                    st.session_state.rag_chain = chain
                    st.session_state.current_file = file_id
                    st.session_state.messages = [] # Reset chat on new file
                    st.success("Document processed! You can now start chatting.")

# Chat Interface
st.divider()

# Display chat messages from history
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# React to user input
if prompt := st.chat_input("Ask a question about your document..."):
    # Check if chain is ready
    if not st.session_state.rag_chain:
        st.error("Please upload a document first.")
    else:
        # 1. Display user message
        st.chat_message("user").markdown(prompt)
        st.session_state.messages.append({"role": "user", "content": prompt})

        # 2. Generate response
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                try:
                    response = st.session_state.rag_chain.invoke({"input": prompt})
                    answer = response["answer"]
                    st.markdown(answer)
                    
                    # 3. Add assistant response to history
                    st.session_state.messages.append({"role": "assistant", "content": answer})
                except Exception as e:
                    st.error(f"An error occurred: {e}")